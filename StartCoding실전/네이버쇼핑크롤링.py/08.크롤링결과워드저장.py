from ast import keyword
from turtle import title
import pyautogui
import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from docx.shared import Inches

keyWord = pyautogui.prompt("검색어를 입력하시오>>>")
lastPage = int(pyautogui.prompt("몇 페이지까지 크롤링 할까요?"))

#1.워드 생성
doc = Document()

pageNum = 1
for i in range(1, int(lastPage)*10, 10):
    print(f"{pageNum}페이지 크롤링 중입니다.=================")
    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyWord}&start={i}")
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')
    articles = soup.select("div.info_group") #뉴스기사 10개 추출
    for article in articles:
        links = article.select("a.info")    # a태그인데 info class인 애들
        if len(links) >= 2: #링크가 2개 이상이면
            url = links[1].attrs['href']    #두 번째 링크의 URL 추출
            response = requests.get(url, headers={'User-agent' : 'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            print(url)
            #만약 스포츠뉴스 기사면
            if "sports" in response.url:
                title = soup.select_one("title")
                content = soup.select_one("#newsEndContents")
                divs = content.select("div")
                for div in divs:
                    div.decompose()
                paragraphs = content.select("p")
                for p in paragraphs:
                    p.decompose()
            #만약 연예뉴스 기사면
            elif "entertain" in response.url:
                title = soup.select_one("h4.title")
                content = soup.select_one("#articeBody")
            else:
                title = soup.select_one(".media_end_head_headline")
                content = soup.select_one("#dic_area")
            print(url)
            print(title)
            print(content.text.strip())
            time.sleep(0.3)
            #2.워드 데이터 생성
            doc.add_heading(title, level=0)
            doc.add_paragraph(url)
            doc.add_paragraph(content.text.strip())
    pageNum = pageNum + 1
#3. 워드 저장
doc.save(f"{keyWord}_result.docx")