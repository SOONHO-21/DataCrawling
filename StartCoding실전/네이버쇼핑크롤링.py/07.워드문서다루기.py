from docx import Document
from docx.shared import Inches

#1.워드 생성
doc = Document()

#2.워드 데이터 생성
doc.add_heading('기사 제목', level=0)
doc.add_paragraph('기사 링크')
doc.add_paragraph('기사 본문')

#3. 워드 저장
doc.save('demo.docx')